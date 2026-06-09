"""
订单解析工具 - 将客户订单（任意格式）解析为原始结构化JSON

支持 Excel/图片/PDF/Word/文字，通过LLM理解内容，输出原始JSON（多门店支持）

设计原则：
- 任意格式 → LLM解析 → 原始JSON（保持原始字段名）
- 规则库（field_transformer）负责字段名标准化
- 两者互补：LLM做理解，规则库做校正
"""
import os
import re
import json
from typing import Optional, Dict, Any

# LLM Router 导入
try:
    from learn.llm import LLMRouter
except ImportError:
    LLMRouter = None  # 兼容：未安装 LLM 模块时降级到旧逻辑


# 全局 Router 实例（延迟初始化）
_llm_router = None


def _get_llm_router():
    """获取 LLM Router 实例（延迟初始化）"""
    global _llm_router
    if _llm_router is None:
        if LLMRouter is None:
            raise RuntimeError("LLM Router 模块未安装，无法调用 LLM")
        # 配置文件路径：相对于 skill 目录
        skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        config_path = os.path.join(skill_dir, "config", "llm.yaml")
        _llm_router = LLMRouter(config_path)
    return _llm_router


# =============================================================================
# 公共入口函数
# =============================================================================

def parse(order_input: str, order_type: str = "auto",
         db_config: Optional[Dict] = None,
         ocr_result: Dict = None,
         pdf_text: str = None,
         docx_text: str = None) -> Dict[str, Any]:
    """
    解析任意格式客户订单的主入口

    Args:
        order_input: 文件路径 或 文本内容
        order_type: auto/excel/image/pdf/word/text
        db_config: 数据库配置（当前未使用，保留参数兼容）
        ocr_result: 图片OCR结果（当图片解析需要外部工具识别时，由调用方调用image工具后传入）
                    结构: {"text": "识别出的文本内容", ...}
        pdf_text: PDF文本提取结果（当PDF需要外部工具提取时，由调用方调用pdf工具后传入）
        docx_text: Word文本提取结果（当Word需要外部工具提取时，由调用方手动提取后传入）

    Returns:
        {
            "success": bool,
            "stores": {...},      # 多门店时
            "store_name": str,    # 单门店时（兼容）
            "items": [...],
            "warnings": [...],
            "extracted_from": str,
            "_parse_method": str,
        }
        
    特殊返回（需要外部工具处理）：
        {"success": False, "need_ocr": True, "image_path": str, ...}
        {"success": False, "need_pdf": True, "pdf_path": str, ...}
        {"success": False, "need_docx": True, "docx_path": str, ...}
    """
    # 自动检测类型
    if order_type == "auto":
        order_type = detect_input_type(order_input)

    # 移除代理（防止LLM调用失败）
    for k in ["SOCKS_PROXY", "ALL_PROXY", "HTTPS_PROXY", "HTTP_PROXY"]:
        os.environ.pop(k, None)
    os.environ.pop("no_proxy", None)
    os.environ.pop("NO_PROXY", None)

    # 根据类型选择解析路径
    if order_type == "excel":
        return _parse_excel(order_input)
    elif order_type == "image":
        # 优先使用外部工具的OCR结果
        if ocr_result:
            return parse_with_ocr_result(order_input, ocr_result)
        return _parse_image(order_input)
    elif order_type == "pdf":
        # 优先使用外部工具提取的PDF文本
        if pdf_text:
            return parse_with_pdf_text_result(order_input, pdf_text)
        return _parse_pdf(order_input)
    elif order_type == "word":
        # 优先使用外部工具提取的Word文本
        if docx_text:
            return parse_with_docx_text_result(order_input, docx_text)
        return _parse_word(order_input)
    elif order_type == "text":
        return _parse_text(order_input)
    else:
        return {"success": False, "error": f"不支持的格式: {order_type}"}


def parse_with_ocr_result(image_path: str, ocr_result: Dict) -> Dict[str, Any]:
    """
    使用外部工具(image工具)OCR结果解析图片订单
    
    Args:
        image_path: 图片路径
        ocr_result: OCR工具返回的结构化结果
                   期望格式: {"text": "识别出的文本", ...} 或直接是包含文本的dict
    
    Returns:
        LLM解析后的订单结构化数据
    """
    # 从ocr_result中提取文本
    ocr_text = None
    if isinstance(ocr_result, dict):
        ocr_text = ocr_result.get("text", "")
    
    # 兜底：如果dict中没有text字段或其值为空，尝试从其他常见字段找
    if not ocr_text or str(ocr_text).strip() in ("", "None"):
        if isinstance(ocr_result, dict):
            for k in ["content", "description", "raw_text", "full_text"]:
                if k in ocr_result and ocr_result[k]:
                    ocr_text = str(ocr_result[k])
                    break
    
    # 如果仍然没有有效文本，返回错误
    if not ocr_text or str(ocr_text).strip() in ("", "None"):
        return {
            "success": False,
            "error": "OCR结果为空，请检查image工具返回的内容",
            "extracted_from": "image",
        }
    
    return _parse_image_with_ocr(image_path, str(ocr_text))


def parse_with_pdf_text_result(pdf_path: str, pdf_text: str) -> Dict[str, Any]:
    """使用外部工具(pdf工具)提取的文本解析PDF订单"""
    if not pdf_text or not pdf_text.strip():
        return {"success": False, "error": "PDF文本提取结果为空", "extracted_from": "pdf"}
    return _parse_pdf_with_text(pdf_path, pdf_text)


def parse_with_docx_text_result(docx_path: str, docx_text: str) -> Dict[str, Any]:
    """使用外部工具提取的Word文本解析Word订单"""
    if not docx_text or not docx_text.strip():
        return {"success": False, "error": "Word文档文本提取结果为空", "extracted_from": "word"}
    return _parse_word_with_text(docx_path, docx_text)


def detect_input_type(order_input: str) -> str:
    """自动检测输入类型: excel/image/pdf/word/text"""
    if not order_input:
        return "text"
    if os.path.exists(order_input):
        ext = os.path.splitext(order_input)[1].lower()
        if ext in [".xlsx", ".xls"]:
            return "excel"
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp"]:
            return "image"
        elif ext == ".pdf":
            return "pdf"
        elif ext in [".docx", ".doc"]:
            return "word"
        return "image"
    s = str(order_input).strip()
    if s.endswith((".xlsx", ".xls")):
        return "excel"
    elif s.endswith((".jpg", ".png", ".jpeg", ".webp")):
        return "image"
    elif s.endswith(".pdf"):
        return "pdf"
    if "\n" in s or any(kw in s for kw in ["订单号", "门店", "商品", "数量", "规格"]):
        return "text"
    return "text"


# =============================================================================
# Excel解析
# =============================================================================

def _excel_to_text(file_path: str, max_columns: int = None) -> str:
    """将 Excel 文件转为文本描述
    
    读取所有列（包括收货人姓名、手机号、收货地址等扩展字段）
    """
    import pandas as pd
    df = pd.read_excel(file_path, header=None)
    lines = ["【Excel订单内容开始】"]
    n_cols = max_columns if max_columns else len(df.columns)
    for i in range(len(df)):
        row_data = []
        for j in range(min(n_cols, len(df.columns))):
            val = df.iloc[i, j]
            if val is not None and str(val) != 'nan' and str(val).strip():
                row_data.append(str(val).strip())
        if row_data:
            lines.append(f"行{i + 1}: {' | '.join(row_data)}")
    lines.append("【Excel订单内容结束】")
    return "\n".join(lines)

def _parse_excel(file_path: str) -> Dict[str, Any]:
    """Excel 解析：LLM + fallback正则"""
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}"}
    try:
        text_content = _excel_to_text(file_path)
        result = _call_llm(text_content, "excel")
        if result.get("success"):
            result["extracted_from"] = "excel"
            result["_parse_method"] = "llm"
        else:
            result = _parse_excel_regex(file_path)
            result["extracted_from"] = "excel"
            result["_parse_method"] = "regex_fallback"
        return result
    except Exception as e:
        return {"success": False, "error": f"Excel解析失败: {str(e)}"}


def _parse_image(file_path: str) -> Dict[str, Any]:
    """图片解析：返回需要外部OCR工具的标记"""
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}"}
    return {
        "success": False,
        "need_ocr": True,
        "image_path": file_path,
        "message": "图片需要OCR识别，请使用 image 工具识别后回传",
        "extracted_from": "image",
    }


def _parse_image_with_ocr(image_path: str, ocr_text: str) -> Dict[str, Any]:
    """图片解析（已有OCR文本结果）"""
    try:
        result = _call_llm(ocr_text, "text")
        if result.get("success"):
            result["extracted_from"] = "image"
            result["_parse_method"] = "llm_ocr"
        return result
    except Exception as e:
        return {"success": False, "error": f"图片解析失败: {str(e)}"}


def _parse_pdf(file_path: str) -> Dict[str, Any]:
    """PDF解析：返回需要外部工具的标记"""
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}"}
    return {
        "success": False,
        "need_pdf": True,
        "pdf_path": file_path,
        "message": "PDF需要文本提取，请使用 pdf 工具处理后回传",
        "extracted_from": "pdf",
    }


def _parse_pdf_with_text(pdf_path: str, pdf_text: str) -> Dict[str, Any]:
    """PDF解析（已有提取文本）"""
    try:
        result = _call_llm(pdf_text, "text")
        if result.get("success"):
            result["extracted_from"] = "pdf"
            result["_parse_method"] = "llm_pdf"
        return result
    except Exception as e:
        return {"success": False, "error": f"PDF解析失败: {str(e)}"}


def _parse_word(file_path: str) -> Dict[str, Any]:
    """Word解析：尝试用python-docx提取文本，失败则返回need_docx标记"""
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}"}
    try:
        from docx import Document
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    full_text.append(row_text)
        docx_text = "\n".join(full_text)
        if docx_text.strip():
            return _parse_word_with_text(file_path, docx_text)
        else:
            return {
                "success": False,
                "need_docx": True,
                "docx_path": file_path,
                "message": "Word文档内容为空，请手动提取文本后回传",
                "extracted_from": "word",
            }
    except ImportError:
        return {
            "success": False,
            "need_docx": True,
            "docx_path": file_path,
            "message": "python-docx库未安装，无法自动提取Word文档",
            "extracted_from": "word",
        }
    except Exception as e:
        return {
            "success": False,
            "need_docx": True,
            "docx_path": file_path,
            "message": f"Word文档解析失败: {str(e)}，请手动提取文本后回传",
            "extracted_from": "word",
        }


def _parse_word_with_text(docx_path: str, docx_text: str) -> Dict[str, Any]:
    """Word解析（已有提取文本）"""
    try:
        result = _call_llm(docx_text, "text")
        if result.get("success"):
            result["extracted_from"] = "word"
            result["_parse_method"] = "llm_docx"
        return result
    except Exception as e:
        return {"success": False, "error": f"Word解析失败: {str(e)}"}


def _parse_text(text: str) -> Dict[str, Any]:
    """纯文本解析：LLM"""
    if not text or not text.strip():
        return {"success": False, "error": "文本内容为空"}
    try:
        result = _call_llm(text, "text")
        if result.get("success"):
            result["extracted_from"] = "text"
            result["_parse_method"] = "llm"
        return result
    except Exception as e:
        return {"success": False, "error": f"文本解析失败: {str(e)}"}


# =============================================================================
# Excel正则fallback
# =============================================================================

def _parse_excel_regex(file_path: str) -> Dict[str, Any]:
    """Excel正则fallback（LLM失败时用）"""
    try:
        import pandas as pd
        df_raw = pd.read_excel(file_path, header=None)
        header_detail_result = _parse_excel_header_detail(df_raw)
        if header_detail_result.get("success"):
            return header_detail_result

        data_start = 6
        col_mapping = {}
        for row_idx in [0, 1, 2, 3, 4, 5, 6]:
            if row_idx >= len(df_raw):
                continue
            row = df_raw.iloc[row_idx]
            temp_map = {}
            for col_idx, cell in enumerate(row):
                cell_str = str(cell).strip() if cell else ""
                std_field = _find_standard_field(cell_str)
                if std_field:
                    temp_map[col_idx] = std_field
            if len(temp_map) > len(col_mapping):
                col_mapping = temp_map
                data_start = row_idx + 1
        field_to_col = {v: k for k, v in col_mapping.items()}
        stores_data = {}
        for row_idx in range(data_start, len(df_raw)):
            if row_idx >= len(df_raw):
                break
            col_store = field_to_col.get("store_name", 1)
            store_name = str(df_raw.iloc[row_idx, col_store]).strip() if pd.notna(df_raw.iloc[row_idx, col_store]) else ""
            for prefix in ['河北-', '天津-', '沧州-', '创宇-', '盐城创宇-']:
                if store_name.startswith(prefix):
                    store_name = store_name[len(prefix):]
                    break
            if not store_name or store_name in ["nan", "门店名称", "往来单位名称"]:
                continue
            col_pn = field_to_col.get("product_name", 4)
            col_qty = field_to_col.get("quantity", 7)
            col_unit = field_to_col.get("unit", 6)
            col_spec = field_to_col.get("spec", 5)
            product_name = str(df_raw.iloc[row_idx, col_pn]).strip() if pd.notna(df_raw.iloc[row_idx, col_pn]) else ""
            if not product_name or product_name in ["nan", "商品名称", "品名"]:
                continue
            qty_raw = df_raw.iloc[row_idx, col_qty]
            quantity = int(qty_raw) if pd.notna(qty_raw) and isinstance(qty_raw, (int, float)) else 0
            unit = str(df_raw.iloc[row_idx, col_unit]).strip() if pd.notna(df_raw.iloc[row_idx, col_unit]) else "件"
            spec = str(df_raw.iloc[row_idx, col_spec]).strip() if pd.notna(df_raw.iloc[row_idx, col_spec]) else ""
            if store_name not in stores_data:
                stores_data[store_name] = {"store_name": store_name, "order_no": "", "items": []}
            stores_data[store_name]["items"].append({
                "seq": len(stores_data[store_name]["items"]) + 1,
                "product_name": product_name, "spec": spec, "quantity": quantity,
                "unit": unit, "remark": ""
            })
        return {
            "success": True, "confidence": 0.5, "stores": stores_data,
            "warnings": ["正则fallback解析，置信度较低"],
        }
    except Exception as e:
        return {"success": False, "error": f"Excel正则解析失败: {str(e)}"}


def _parse_excel_header_detail(df_raw) -> Dict[str, Any]:
    """
    解析常见的「订单头 + 商品明细表」Excel。

    示例：
      订单编号 | DH-O-...
      公司名称 | 任丘三中店 | ... | 联系人 | 任建华
      收货地址 | ...       | ... | 联系电话 | 181...
      商 品 明 细
      序号 | 商品编码 | 商品名称 | 商品规格 | 数量 | 计量单位 | 备注
      1    | P...     | 椰果果粒 | ...      | 3    | 件
    """
    import pandas as pd

    def cell(row_idx: int, col_idx: int) -> str:
        if row_idx < 0 or row_idx >= len(df_raw) or col_idx < 0 or col_idx >= len(df_raw.columns):
            return ""
        val = df_raw.iloc[row_idx, col_idx]
        if pd.isna(val):
            return ""
        text = str(val).strip()
        if text.lower() == "nan":
            return ""
        return text

    def normalize_label(value: str) -> str:
        return re.sub(r"[\s\u3000：:]+", "", str(value or "")).strip()

    def row_values(row_idx: int) -> list:
        return [cell(row_idx, col_idx) for col_idx in range(len(df_raw.columns))]

    def find_value_after_label(row_idx: int, labels: list) -> str:
        values = row_values(row_idx)
        normalized_labels = {normalize_label(label) for label in labels}
        for col_idx, value in enumerate(values):
            if normalize_label(value) in normalized_labels:
                for next_col in range(col_idx + 1, len(values)):
                    if values[next_col]:
                        return values[next_col]
        return ""

    order_no = ""
    store_name = ""
    contact_person = ""
    phone = ""
    address = ""
    header_row_idx = None

    for row_idx in range(len(df_raw)):
        values = row_values(row_idx)
        normalized = [normalize_label(v) for v in values]
        non_empty = [v for v in normalized if v]

        if not order_no:
            order_no = find_value_after_label(row_idx, ["订单编号", "订单号", "单号"])
        if not store_name:
            store_name = find_value_after_label(row_idx, ["公司名称", "门店名称", "门店", "客户名称", "店名"])
        if not contact_person:
            contact_person = find_value_after_label(row_idx, ["联系人", "收货人", "收货人姓名", "收件人"])
        if not phone:
            phone = find_value_after_label(row_idx, ["联系电话", "手机号", "手机", "电话"])
        if not address:
            address = find_value_after_label(row_idx, ["收货地址", "地址", "配送地址"])

        has_item_header = (
            any(v in {"序号", "商品编码", "商品名称", "品名"} for v in normalized)
            and any(v in {"数量", "订货数量", "订单数量"} for v in normalized)
            and any(v in {"商品名称", "品名", "货品名称"} for v in normalized)
        )
        if has_item_header:
            header_row_idx = row_idx
            break

        if any("商品明细" in v or "商品詳情" in v for v in non_empty):
            continue

    if header_row_idx is None or not store_name:
        return {"success": False, "error": "未识别到订单头+商品明细格式"}

    header_mapping = {}
    for col_idx, value in enumerate(row_values(header_row_idx)):
        label = normalize_label(value)
        if not label:
            continue
        std_field = _find_standard_field(label)
        if std_field:
            header_mapping[std_field] = col_idx
        elif label in {"计量单位", "销售单位"}:
            header_mapping["unit"] = col_idx
        elif label in {"备注", "备注信息"}:
            header_mapping["remark"] = col_idx

    if "product_name" not in header_mapping or "quantity" not in header_mapping:
        return {"success": False, "error": "商品明细表头缺少商品名称或数量"}

    items = []
    for row_idx in range(header_row_idx + 1, len(df_raw)):
        first_cell = normalize_label(cell(row_idx, 0))
        row_text = "".join(row_values(row_idx))
        if not row_text:
            continue
        if first_cell.startswith("合计") or first_cell in {"合计", "总计"}:
            break

        product_name = cell(row_idx, header_mapping["product_name"])
        if not product_name or normalize_label(product_name) in {"商品名称", "品名"}:
            continue

        quantity = _safe_int_excel(cell(row_idx, header_mapping["quantity"]))
        if quantity <= 0:
            continue

        product_code = cell(row_idx, header_mapping.get("product_code", -1))
        spec = cell(row_idx, header_mapping.get("spec", -1))
        unit = cell(row_idx, header_mapping.get("unit", -1)) or "件"
        remark = cell(row_idx, header_mapping.get("remark", -1))

        seq_raw = cell(row_idx, header_mapping.get("seq", 0))
        seq = _safe_int_excel(seq_raw) or len(items) + 1

        items.append({
            "seq": seq,
            "product_code": product_code,
            "product_name": product_name,
            "spec": spec,
            "quantity": quantity,
            "unit": unit,
            "remark": remark,
        })

    if not items:
        return {"success": False, "error": "未识别到商品明细"}

    for prefix in ['河北-', '天津-', '沧州-', '创宇-', '盐城创宇-']:
        if store_name.startswith(prefix):
            store_name = store_name[len(prefix):]
            break

    return {
        "success": True,
        "confidence": 0.8,
        "stores": {
            store_name: {
                "store_name": store_name,
                "contact_person": contact_person,
                "phone": phone,
                "address": address,
                "order_no": order_no,
                "items": items,
            }
        },
        "warnings": ["Excel订单头+商品明细fallback解析"],
    }


def _safe_int_excel(value) -> int:
    """Excel 数量/序号转 int，支持 '3.0'、'数量：3' 等输入。"""
    if value is None:
        return 0
    text = str(value).strip()
    if not text or text.lower() == "nan":
        return 0
    match = re.search(r"\d+(?:\.\d+)?", text)
    if not match:
        return 0
    return int(float(match.group(0)))


def _find_standard_field(input_name: str) -> Optional[str]:
    """将任意字段名映射到标准字段名"""
    FIELD_MAP = {
        "store_name": ["往来单位名称", "门店名称", "门店", "店名", "店铺", "收货方"],
        "seq": ["序号", "行号"],
        "product_name": ["商品名称", "商品名", "品名", "商品", "货品名称", "产品名称"],
        "quantity": ["数量", "件数", "箱数", "qty", "出库数量", "订货数量"],
        "unit": ["销售单位", "单位", "件", "箱", "包装单位"],
        "spec": ["规格", "包装规格", "规格型号", "商品规格", "产品规格"],
        "product_code": ["商品编码", "编码", "货号", "SKU", "商品SKU"],
        "order_no": ["单据日期", "订单号", "单号", "订单编号"],
    }
    input_lower = input_name.lower().strip()
    for std_field, aliases in FIELD_MAP.items():
        if input_lower in [a.lower() for a in aliases]:
            return std_field
    return None


# =============================================================================


_SINGLE_PROMPT_TEMPLATE = (
    "你是一个订单数据提取专家。请从客户订单中提取结构化信息。\n"
    "\n"
    "【重要】一个订单可能包含多个门店，每个门店有独立的商品明细。请将所有门店都提取出来。\n"
    "\n"
    "【提取规则】\n"
    "1. 仔细辨认订单中的门店名称（收货方/店铺名）\n"
    "2. 每个门店独立列出，包含该门店所有商品\n"
    "3. 每个商品提取：seq(序号)、product_name(商品名称)、spec(规格)、quantity(数量)、unit(单位)\n"
    "\n"
    "【输出格式】直接输出JSON（不要有其他文字）：\n"
    "{{\n"
    "  \"confidence\": 0.95,\n"
    "  \"stores\": {{\n"
    "    \"门店A名称\": {{\n"
    "      \"store_name\": \"门店A完整名称\",\n"
    "      \"contact_person\": \"收货人姓名\",\n"
    "      \"phone\": \"联系电话\",\n"
    "      \"address\": \"收货地址\",\n"
    "      \"order_no\": \"订单编号（如有）\",\n"
    "      \"items\": [\n"
    "        {{ \"seq\": 1, \"product_name\": \"商品名称\", \"spec\": \"规格（如有）\", \"quantity\": 数量, \"unit\": \"单位\", \"product_code\": \"商品编码（如有）\", \"remark\": \"备注（如有）\" }}\n"
    "      ]\n"
    "    }}\n"
    "  }},\n"
    "  \"warnings\": [\"告警信息（如有）\"]\n"
    "}}\n"
    "\n"
    "【规则】\n"
    "1. 一个订单文件如果包含多个门店，必须把所有门店都列在stores对象里\n"
    "2. 每个门店的items只包含该门店的商品\n"
    "3. product_name必须是完整商品名称，quantity必须是数字类型\n"
    "4. 如果某些字段在订单中找不到，直接省略该字段，不要填null或空字符串\n"
    "\n"
    "【订单内容】：\n"
    "{content}"
)

# LLM 调用
# =============================================================================

_MULTI_PROMPT_TEMPLATE = (
    "你是一个订单数据提取专家。请从客户订单中提取结构化信息。\n"
    "\n"
    "【重要】一个订单可能包含多个门店，每个门店有独立的商品明细。请将所有门店都提取出来。\n"
    "\n"
    "【订单格式说明】\n"
    "Excel订单中，每隔若干行有一个「主行」（包含订单编号），该行的收货人姓名、联系电话、收货地址列可能有值。\n"
    "后续跟的是该订单的商品明细行（无订单编号）。\n"
    "示例结构：\n"
    "  主行: 订单编号 | 支付单号 | ... | 收货人姓名 | 收货人手机号 | 收货地址\n"
    "  商品行: [空] | [空] | [空] | [空] | [空] | [空] | 商品名称 | 商品规格 | 商品数量\n"
    "\n"
    "【提取规则】\n"
    "1. 遇到包含订单编号的主行时：\n"
    "   - 这是新门店/新订单的开始\n"
    "   - 从该行提取 contact_person（收货人姓名）、phone（收货人手机号）、address（收货地址）\n"
    "   - 同一订单后续商品行共用这些联系信息\n"
    "2. 每个门店的contact_person/phone/address只需提取一次（主行优先）\n"
    "\n"
    "【输出格式】直接输出JSON（不要有其他文字）：\n"
    "{{\n"
    "  \"confidence\": 0.95,\n"
    "  \"stores\": {{\n"
    "    \"门店A名称\": {{\n"
    "      \"store_name\": \"门店A完整名称\",\n"
    "      \"contact_person\": \"收货人姓名\",\n"
    "      \"phone\": \"联系电话\",\n"
    "      \"address\": \"收货地址\",\n"
    "      \"order_no\": \"订单编号（如有）\",\n"
    "      \"items\": [\n"
    "        {{ \"seq\": 1, \"product_name\": \"商品名称\", \"spec\": \"规格（如有）\", \"quantity\": 数量, \"unit\": \"单位\", \"product_code\": \"商品编码（如有）\", \"remark\": \"备注（如有）\" }}\n"
    "      ]\n"
    "    }}\n"
    "  }},\n"
    "  \"warnings\": [\"告警信息（如有）\"]\n"
    "}}\n"
    "\n"
    "【规则】\n"
    "1. 一个订单文件如果包含多个门店，必须把所有门店都列在stores对象里\n"
    "2. 每个门店的items只包含该门店的商品\n"
    "3. product_name必须是完整商品名称，quantity必须是数字类型\n"
    "4. 如果某些字段（如contact_person、phone、address）在订单中找不到，直接省略该字段，不要填null或空字符串\n"
    "\n"
    "【订单内容】：\n"
    "{content}"
)


# =============================================================================
# LLM 调用
# =============================================================================

def _load_env():
    """加载 .env 环境变量（如果存在）"""
    env_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "..", "..", ".env")
    if os.path.exists(env_path):
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                if k.strip() not in os.environ:
                    os.environ[k.strip()] = v.strip()


def _call_llm(text_content: str, content_type: str = "text") -> Dict[str, Any]:
    """
    调用 LLM 解析订单文本
    
    使用 LLM Router，支持多种 provider（默认 openclaw 平台内嵌模型）
    支持故障回退链
    """
    try:
        # 加载 .env
        _load_env()
        
        # 构建 prompt
        prompt_template = _MULTI_PROMPT_TEMPLATE if content_type == "excel" else _SINGLE_PROMPT_TEMPLATE
        prompt = prompt_template.format(
            content=text_content[:12000] if len(text_content) > 12000 else text_content
        )
        
        system_msg = "你是一个专业的订单数据提取专家。请从用户提供的订单内容中准确提取结构化信息。"
        
        # 通过 Router 调用 LLM
        router = _get_llm_router()
        response = router.chat(system=system_msg, user=prompt)
        
        raw_output = response.text
        
        # 解析 JSON（与原来逻辑一致）
        if raw_output.startswith("```"):
            raw_output = raw_output.split("```")[1]
            if raw_output.startswith("json"):
                raw_output = raw_output[4:]
            raw_output = raw_output.strip().strip("```")
        
        # 移除思考过程标签
        raw_output = re.sub(r'', '', raw_output, flags=re.DOTALL)
        raw_output = re.sub(r'<!--.*?-->', '', raw_output, flags=re.DOTALL)
        raw_output = raw_output.strip()
        
        return _parse_llm_raw(raw_output)
        
    except Exception as e:
        return {"success": False, "error": f"LLM调用失败: {str(e)}"}
def _parse_llm_raw(raw_output: str) -> Dict[str, Any]:
    """解析 LLM 原始输出"""
    code_blocks = raw_output.split("```")
    parsed = None
    for block in code_blocks:
        block = block.strip()
        if block.startswith("json"):
            block = block[4:].strip()
        if not block:
            continue
        try:
            parsed = json.loads(block)
            break
        except (json.JSONDecodeError, ValueError):
            try:
                parsed = json.loads(block.strip())
                break
            except (json.JSONDecodeError, ValueError):
                continue
    if parsed is None:
        text_no_think = raw_output
        think_close_idx = text_no_think.find("</think>")
        if think_close_idx >= 0:
            text_no_think = text_no_think[think_close_idx + 8:]
        text_no_think = re.sub(r'<think>.*?</think>', '', text_no_think, flags=re.DOTALL)
        text_no_think = re.sub(r'<[^>]+>', '', text_no_think)
        text_no_think = text_no_think.strip()
        if text_no_think:
            try:
                parsed = json.loads(text_no_think)
            except (json.JSONDecodeError, ValueError):
                pass
    if parsed is None:
        return {
            "success": False,
            "error": "LLM输出无法解析为JSON",
            "llm_raw_output": raw_output,
        }
    llm_data = parsed
    if "stores" in llm_data:
        stores = {}
        for store_key, store_data in llm_data.get("stores", {}).items():
            store_name = store_data.get("store_name", store_key)
            for prefix in ['河北-', '天津-', '沧州-', '创宇-', '盐城创宇-']:
                if store_name.startswith(prefix):
                    store_name = store_name[len(prefix):]
                    break
            stores[store_name] = store_data
            stores[store_name]["store_name"] = store_name
        return {
            "success": True,
            "confidence": llm_data.get("confidence", 0.9),
            "stores": stores,
            "order_no": llm_data.get("order_no", ""),
            "warnings": llm_data.get("warnings", []),
            "llm_raw_output": raw_output,
        }
    store_info = llm_data.get("store_info", {})
    order_info = llm_data.get("order_info", {})
    items = llm_data.get("items", [])
    store_name = store_info.get("name", "")
    for prefix in ['河北-', '天津-', '沧州-', '创宇-', '盐城创宇-']:
        if store_name.startswith(prefix):
            store_name = store_name[len(prefix):]
            break
    stores = {store_name: {
        "store_name": store_name,
        "contact_person": store_info.get("contact_person", ""),
        "phone": store_info.get("phone", ""),
        "address": store_info.get("address", ""),
        "order_no": order_info.get("order_no", ""),
        "shipper_name": order_info.get("shipper_name", ""),
        "warehouse_name": order_info.get("warehouse_name", ""),
        "items": items,
    }} if store_name else {}
    return {
        "success": True,
        "confidence": llm_data.get("confidence", 0.9),
        "store_name": store_name,
        "stores": stores,
        "items": items,
        "warnings": llm_data.get("warnings", []),
        "llm_raw_output": raw_output,
    }
